import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    ns = nsdecls("w")
    shd = parse_xml(f'<w:shd {ns} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    ns = nsdecls("w")
    tcMar = parse_xml(f'<w:tcMar {ns}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    cell._tc.get_or_add_tcPr().append(tcMar)

def style_table(table, header_bg="1F497D", stripe_bg="F2F4F7"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                set_cell_background(cell, header_bg)
                set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.bold = True
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(9.0)
                        r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if i % 2 == 0:
                    set_cell_background(cell, stripe_bg)
                set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(8.5)

def parse_markdown_to_docx(md_path, docx_path, insert_plots=False):
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    lines = md_text.split("\n")
    i = 0
    in_table = False
    table_lines = []

    def flush_table(t_lines):
        if not t_lines:
            return
        parsed_rows = []
        for line in t_lines:
            if re.match(r'^\s*\|?\s*[-:]+[-| :]*$', line):
                continue
            cols = [c.strip() for c in line.strip().strip('|').split('|')]
            if cols and any(cols):
                parsed_rows.append(cols)

        if not parsed_rows:
            return

        num_cols = max(len(r) for r in parsed_rows)
        # Pad shorter rows
        for r in parsed_rows:
            while len(r) < num_cols:
                r.append("")

        table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
        for r_idx, r_data in enumerate(parsed_rows):
            for c_idx, val in enumerate(r_data):
                clean_val = val.replace("**", "").replace("*", "").replace("`", "")
                cell = table.cell(r_idx, c_idx)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 or len(clean_val) < 15 else WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(clean_val)
                r.font.name = "Times New Roman"
                if r_idx == 0 or "**" in val or "100.0" in val:
                    r.bold = True
        style_table(table)

        # Space after table
        p_after = doc.add_paragraph()
        p_after.paragraph_format.space_after = Pt(6)

    def add_plot_figure(plot_path, caption):
        if os.path.exists(plot_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(10)
            p_img.paragraph_format.space_after = Pt(2)
            doc.add_picture(plot_path, width=Inches(5.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            r = p_cap.add_run(caption)
            r.italic = True
            r.font.size = Pt(9)
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(80, 80, 80)

    in_code_block = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            # Render code lines cleanly
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run(line)
            r.font.name = "Courier New"
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(60, 60, 60)
            i += 1
            continue

        # Check for Markdown Table lines
        if "|" in stripped and not stripped.startswith("#"):
            table_lines.append(line)
            i += 1
            continue
        else:
            if table_lines:
                flush_table(table_lines)
                table_lines = []

        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped[2:])
            r.bold = True
            r.font.size = Pt(17)
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(31, 73, 125)
        elif stripped.startswith("## "):
            heading_text = stripped[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(heading_text)
            r.bold = True
            r.font.size = Pt(13)
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(31, 73, 125)
        elif stripped.startswith("### "):
            subheading_text = stripped[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(subheading_text)
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(40, 40, 40)
        elif stripped.startswith("---"):
            i += 1
            continue
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            
            # Format bold runs within bullets
            parts = re.split(r'(\*\*.*?\*\*)', stripped[2:])
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)
        elif stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. ") or stripped.startswith("4. ") or stripped.startswith("5. "):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)
        else:
            # Paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.15

            # If it is author info
            if "Arpit Panigrahi" in stripped or "Vellore" in stripped or "Email" in stripped or "GitHub" in stripped:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    r = p.add_run(part[1:-1])
                    r.italic = True
                else:
                    r = p.add_run(part)
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)

        # Insert relevant plots dynamically under specific sections in IEEE paper
        if insert_plots:
            if "### B. The 52% Unconstrained Legality Ceiling" in stripped:
                add_plot_figure("reports/figures/plot1_legality_rate.png", "Figure 1: Move Legality Rate Across Experimental Conditions. Unconstrained models fail at 51.8%–54.5%, while Quoted DMC achieves 100.0% zero-hallucination compliance.")
            elif "### D. Granular Latency Profiling" in stripped:
                add_plot_figure("reports/figures/plot2_latency_decomposition.png", "Figure 2: Latency Decomposition: Initial Cold-Start Disk Load vs. Steady-State Warm Inference vs. p95 Tail Ceiling.")
            elif "## VI. Tokenizer & Representation Ablation Study" in stripped:
                add_plot_figure("reports/figures/plot3_acpl_quality.png", "Figure 3: Stockfish 18 Average Centipawn Loss (ACPL). Constrained play sustains authentic club-level quality (55.8–67.0 cp) across full multi-turn games.")
                add_plot_figure("reports/figures/plot4_token_compression.png", "Figure 4: Prompt Token Footprint Comparison. Dynamic Move Compression cuts prefill compute payload by 45%.")

        i += 1

    if table_lines:
        flush_table(table_lines)

    doc.save(docx_path)
    print(f"✓ Successfully compiled 100% full content of {md_path} -> {docx_path}")

if __name__ == "__main__":
    parse_markdown_to_docx("docs/research_paper_ieee.md", "docs/research_paper_ieee.docx", insert_plots=True)
    parse_markdown_to_docx("docs/comparative_literature_study.md", "docs/comparative_literature_study.docx", insert_plots=False)
