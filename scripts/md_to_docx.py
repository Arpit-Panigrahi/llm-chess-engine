#!/usr/bin/env python3
"""
md_to_docx.py — Convert research_paper_ieee.md to a formatted DOCX file.
Uses python-docx to produce a professional IEEE-style Word document.
"""

import os
import re
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PLOTS_DIR = os.path.join(PROJECT_ROOT, "reports", "experiment_matrix", "plots")


def set_cell_shading(cell, color_hex):
    """Apply shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def add_styled_paragraph(doc, text, style_name, font_size=None, bold=False, italic=False,
                         alignment=None, space_after=None, space_before=None, color=None):
    """Add a paragraph with specific styling."""
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if font_size:
        run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = 'Times New Roman'
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def add_mixed_paragraph(doc, parts, style_name='Normal', alignment=None, space_after=None):
    """Add paragraph with mixed bold/italic/code runs."""
    p = doc.add_paragraph(style=style_name)
    for text, bold, italic, code in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def parse_inline(text):
    """Parse inline markdown into (text, bold, italic, code) tuples."""
    parts = []
    # Pattern: **bold**, *italic*, `code`
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))'
    for match in re.finditer(pattern, text):
        if match.group(2):  # bold
            parts.append((match.group(2), True, False, False))
        elif match.group(3):  # italic
            parts.append((match.group(3), False, True, False))
        elif match.group(4):  # code
            parts.append((match.group(4), False, False, True))
        elif match.group(5):  # plain
            parts.append((match.group(5), False, False, False))
    return parts


def build_docx():
    doc = Document()

    # ── Page Setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.91)
        section.right_margin = Cm(1.91)

    # ── Default Style ──
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.space_before = Pt(0)

    # ══════════════════════════════════════════════════════════════
    # TITLE
    # ══════════════════════════════════════════════════════════════
    add_styled_paragraph(
        doc,
        "Constraining Large Language Model Chess Move Generation: "
        "A Prompt-Level Legal Move Injection Approach to Eliminating Hallucinations",
        'Normal', font_size=16, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
    )

    # Author
    add_styled_paragraph(
        doc, "Arpit Panigrahi", 'Normal', font_size=12, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2
    )

    # Affiliation
    add_styled_paragraph(
        doc,
        "School of Computer Science and Engineering, "
        "Vellore Institute of Technology (VIT), Chennai, Tamil Nadu, India",
        'Normal', font_size=10, italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2
    )

    # Email
    add_styled_paragraph(
        doc, "Email: arpitpanigrahi06@gmail.com", 'Normal', font_size=10,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2
    )

    # GitHub
    add_styled_paragraph(
        doc,
        "GitHub: https://github.com/Arpit-Panigrahi/llm-chess-engine",
        'Normal', font_size=10,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
    )

    doc.add_paragraph().add_run("─" * 80).font.size = Pt(6)

    # ══════════════════════════════════════════════════════════════
    # ABSTRACT
    # ══════════════════════════════════════════════════════════════
    add_styled_paragraph(doc, "Abstract", 'Normal', font_size=11, bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, space_before=6)

    abstract_text = (
        "Large Language Models (LLMs) demonstrate broad linguistic competence but struggle with "
        "tasks requiring strict adherence to formal rule systems, such as legal chess move generation. "
        "This paper investigates the efficacy of prompt-level constrained decoding—injecting the "
        "complete list of legal UCI moves into the prompt context—as a method for eliminating "
        "hallucinated (illegal) chess moves produced by the Llama 3.1 8B model. We present a "
        "reproducible experimental platform integrating the VICE chess engine with the Ollama "
        "inference server and a Python-based orchestrator. Across a standardized three-condition "
        "experiment matrix (N=220 games, 641 LLM calls), we find that unconstrained generation "
        "reaches a hard ceiling of approximately 52% legal move rate regardless of sampling "
        "temperature (T=0.2: 52.6%, T=0.8: 52.4%), while prompt-level constraint injection achieves "
        "a perfect 100.0% legal move rate. Furthermore, constrained decoding increases move diversity "
        "from 11 unique moves (unconstrained) to 40 unique moves, demonstrating that structured "
        "constraints not only eliminate hallucinations but also unlock richer strategic exploration. "
        "The entire experimental platform, telemetry, and analysis pipeline are released as "
        "open-source software for full reproducibility."
    )
    p = add_styled_paragraph(doc, abstract_text, 'Normal', font_size=9, italic=True,
                             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

    # Keywords
    p = doc.add_paragraph(style='Normal')
    run = p.add_run("Keywords: ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(
        "Large Language Models, Chess, Hallucination, Constrained Decoding, "
        "Prompt Engineering, Llama, UCI Protocol, Move Legality"
    )
    run2.italic = True
    run2.font.size = Pt(9)
    run2.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph().add_run("─" * 80).font.size = Pt(6)

    # ══════════════════════════════════════════════════════════════
    # I. INTRODUCTION
    # ══════════════════════════════════════════════════════════════
    add_styled_paragraph(doc, "I. Introduction", 'Normal', font_size=12, bold=True,
                         space_before=12, space_after=6)

    intro_paras = [
        "Large Language Models (LLMs) have demonstrated remarkable capabilities across diverse "
        "natural language processing tasks, from text generation and summarization to code synthesis "
        "and logical reasoning [1]. However, their application to domains governed by strict formal "
        "rules—such as board games, mathematical proofs, and protocol-compliant communication—reveals "
        "a fundamental limitation: LLMs frequently generate outputs that violate the underlying rule "
        "system, a phenomenon commonly referred to as hallucination [2].",

        "Chess serves as an ideal testbed for studying this limitation. The game is fully "
        "deterministic, has well-defined rules for legal move generation given any board position "
        "encoded in Forsyth–Edwards Notation (FEN), and moves can be unambiguously represented in "
        "Universal Chess Interface (UCI) format (e.g., e2e4, g8f6). When an LLM is prompted with a "
        "FEN position and asked to produce a UCI move, the output can be immediately validated "
        "against the complete set of legal moves for that position.",

        "Prior work has explored LLM chess capabilities through direct prompting [3], fine-tuning "
        "on game databases [4], and integration with classical search algorithms [5]. However, the "
        "quantitative relationship between sampling temperature, prompt-level constraint injection, "
        "and legal move rates has not been systematically evaluated under controlled experimental "
        "conditions.",

        "This paper makes the following contributions:",
    ]

    for para in intro_paras:
        add_styled_paragraph(doc, para, 'Normal', font_size=10,
                             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

    contributions = [
        "An open-source experimental platform integrating the VICE chess engine [6] with the "
        "Ollama inference server [7] and a Python-based orchestrator for automated, reproducible "
        "experiment execution.",
        "A standardized three-condition experiment matrix isolating the effects of sampling "
        "temperature and prompt-level legal move injection on move legality.",
        "Empirical evidence that unconstrained LLM chess move generation reaches a hard ceiling of "
        "approximately 52% legality regardless of temperature, while prompt-level constraint "
        "injection achieves 100% legality with significantly increased move diversity.",
    ]
    for i, c in enumerate(contributions, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f"{c}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

    # ══════════════════════════════════════════════════════════════
    # II. RELATED WORK
    # ══════════════════════════════════════════════════════════════
    add_styled_paragraph(doc, "II. Related Work", 'Normal', font_size=12, bold=True,
                         space_before=12, space_after=6)

    # A.
    add_styled_paragraph(doc, "A. LLM Chess Capabilities", 'Normal', font_size=11, bold=True,
                         italic=True, space_after=4)
    add_styled_paragraph(
        doc,
        "Recent studies have evaluated LLMs on chess tasks, including position evaluation, move "
        "prediction, and full game play. Toshniwal et al. [3] demonstrated that GPT-3.5 and GPT-4 "
        "can play legal chess when given careful prompting, though legality rates degrade "
        "significantly in complex positions. Feng et al. [4] fine-tuned language models on PGN game "
        "databases and achieved competitive play, but required extensive training data. Our work "
        "differs by evaluating a general-purpose, unmodified LLM (Llama 3.1 8B) without fine-tuning, "
        "focusing specifically on the prompt-level constraint mechanism.",
        'Normal', font_size=10, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6
    )

    # B.
    add_styled_paragraph(doc, "B. Hallucination in LLMs", 'Normal', font_size=11, bold=True,
                         italic=True, space_after=4)
    add_styled_paragraph(
        doc,
        "Hallucination—the generation of outputs that are factually incorrect, inconsistent, or "
        "violate domain constraints—is a well-documented challenge in LLM research [2]. In the chess "
        "domain, hallucination manifests as the generation of illegal moves: moves that reference "
        "non-existent squares, move pieces that do not exist at the claimed origin, or violate "
        "movement rules. Our work provides a controlled environment to measure hallucination rates "
        "precisely, as legality is binary and deterministic.",
        'Normal', font_size=10, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6
    )

    # C.
    add_styled_paragraph(doc, "C. Constrained Decoding", 'Normal', font_size=11, bold=True,
                         italic=True, space_after=4)
    add_styled_paragraph(
        doc,
        "Constrained decoding techniques restrict the output space of language models to satisfy "
        "specified constraints. Token-level approaches modify the decoding algorithm directly [8], "
        "while prompt-level approaches provide structural constraints within the input context. Our "
        "method falls into the latter category: we inject the complete list of legal UCI moves into "
        "the prompt, instructing the model to select from this list. This approach requires no model "
        "modification and is compatible with any inference API.",
        'Normal', font_size=10, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6
    )

    # ══════════════════════════════════════════════════════════════
    # III. SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════
    add_styled_paragraph(doc, "III. System Architecture", 'Normal', font_size=12, bold=True,
                         space_before=12, space_after=6)

    add_styled_paragraph(doc, "A. Overview", 'Normal', font_size=11, bold=True,
                         italic=True, space_after=4)

    components = [
        ("VICE Chess Engine (C): ", "A modified version of the open-source VICE engine [6] by "
         "Bluefever Software/Richard Allbert, extended with four custom modules: llm_search.c "
         "(LLM search entry point), http_client.c (Ollama HTTP integration via libcurl), "
         "llm_parser.c (UCI move extraction), and telemetry.c (CSV telemetry logging)."),
        ("Ollama Inference Server: ", "A locally-hosted LLM server [7] running the Llama 3.1 8B "
         "model (4-bit quantization, Q4_0). All inference is performed via the /api/generate REST "
         "endpoint with configurable temperature and seed parameters."),
        ("Python Orchestrator (run_game.py): ", "A pure-Python game runner providing full "
         "programmatic control over temperature, constrained decoding, seed, and model parameters. "
         "White plays random legal moves (seeded); Black plays via LLM."),
        ("Analysis Pipeline (analyze_all.py): ", "An automated discovery, validation, and reporting "
         "tool that scans the runs/ directory, validates data integrity, computes comparative "
         "metrics, and generates plots and a markdown report."),
    ]
    for i, (label, desc) in enumerate(components, 1):
        p = doc.add_paragraph(style='Normal')
        run_label = p.add_run(f"{i}. {label}")
        run_label.bold = True
        run_label.font.name = 'Times New Roman'
        run_label.font.size = Pt(10)
        run_desc = p.add_run(desc)
        run_desc.font.name = 'Times New Roman'
        run_desc.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # B. Robust UCI Parser
    add_styled_paragraph(doc, "B. Robust UCI Parser", 'Normal', font_size=11, bold=True,
                         italic=True, space_before=8, space_after=4)

    add_styled_paragraph(
        doc,
        "The orchestrator includes a multi-stage UCI move parser (extract_uci_move) that separates "
        "formatting variations from true logical chess errors:",
        'Normal', font_size=10, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4
    )

    parser_steps = [
        "Direct UCI Match: Extracts exact 4–5 character UCI patterns (e.g., e2e4, e7e8q).",
        "Long Algebraic Notation (LAN): Strips piece prefixes (e.g., Nb8c6 → b8c6).",
        "Standard Algebraic Notation (SAN): Resolves SAN moves contextually (e.g., Nf6 → g8f6).",
        "Formatting Cleanup: Removes quotes, bolding, hyphens, punctuation.",
    ]
    for i, step in enumerate(parser_steps, 1):
        p = doc.add_paragraph(style='Normal')
        p.add_run(f"  {i}. {step}").font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(2)

    # ══════════════════════════════════════════════════════════════
    # IV. EXPERIMENTAL METHODOLOGY
    # ══════════════════════════════════════════════════════════════
    add_styled_paragraph(doc, "IV. Experimental Methodology", 'Normal', font_size=12, bold=True,
                         space_before=12, space_after=6)

    add_styled_paragraph(doc, "A. Experiment Design", 'Normal', font_size=11, bold=True,
                         italic=True, space_after=4)

    add_styled_paragraph(
        doc,
        "We employ a three-condition within-subject experiment matrix, designed to isolate the "
        "independent effects of (a) sampling temperature and (b) prompt-level legal move injection:",
        'Normal', font_size=10, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6
    )

    # Experiment Design Table
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Condition Tag", "Temperature", "Constrained", "Purpose"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")

    data = [
        ["t02_unconstrained_v2", "0.2", "No", "Low-temperature baseline"],
        ["t08_unconstrained", "0.8", "No", "Mid-temperature baseline"],
        ["t08_constrained", "0.8", "Yes", "Constraint efficacy test"],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # B. Controlled Variables
    add_styled_paragraph(doc, "B. Controlled Variables", 'Normal', font_size=11, bold=True,
                         italic=True, space_after=4)

    controls = [
        "Model: Llama 3.1 8B (via Ollama, 4-bit Q4_0 quantization)",
        "Seed: 42 (passed to both Ollama and the Python random number generator)",
        "Opponent: White plays random legal moves (seeded PRNG: random.Random(seed + game_number))",
        "Turn Cap: 200 ply (100 full moves) maximum per game",
        "Early Termination: Enabled for unconstrained runs (game aborts on first illegal move)",
        "Inference Timeout: 15 seconds per API call",
    ]
    for c in controls:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(c)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    # C. Prompt Design
    add_styled_paragraph(doc, "C. Prompt Design", 'Normal', font_size=11, bold=True,
                         italic=True, space_before=8, space_after=4)

    p = doc.add_paragraph(style='Normal')
    p.add_run("Unconstrained Prompt:").bold = True
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(2)

    unconstrained_prompt = (
        'You are a chess engine playing as Black. The current board FEN is: {fen}. '
        'It is Black\'s turn to move. Respond ONLY with a single UCI move in '
        'source-destination format (e.g., g8f6, e7e5, b8c6, d7d5). Do not include '
        'piece letters, just the two squares. Do not include any other text, '
        'explanations, or formatting.'
    )
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(unconstrained_prompt)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph(style='Normal')
    p.add_run("Constrained Prompt:").bold = True
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(2)

    constrained_prompt = (
        'You are a chess engine playing as Black. The current board FEN is: {fen}. '
        'It is Black\'s turn to move. The ONLY legal moves in this position are: '
        '{legal_moves_json}. You MUST pick exactly one move from that list. Respond '
        'ONLY with a single 4-character UCI move (e.g., e7e5). Do not include any '
        'other text, explanations, or formatting.'
    )
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(constrained_prompt)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(6)

    # D. Metrics
    add_styled_paragraph(doc, "D. Metrics", 'Normal', font_size=11, bold=True,
                         italic=True, space_after=4)
    metrics = [
        "Legal Move Rate: total_legal_moves / total_llm_calls (primary metric)",
        "Unique Moves: Count of distinct UCI moves extracted across all games in a condition",
        "Latency (mean/median): Response time in milliseconds per LLM call",
        "Game Completion: Whether the game reached the turn cap (*) or was aborted",
    ]
    for m in metrics:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(m)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    # E. Hardware
    add_styled_paragraph(doc, "E. Hardware Environment", 'Normal', font_size=11, bold=True,
                         italic=True, space_before=8, space_after=4)
    hw = [
        "OS: Linux (native)",
        "CPU: Multi-core x86_64 processor",
        "Inference: CPU-only Ollama (no GPU acceleration)",
        "Network: Localhost loopback (no network latency)",
    ]
    for h in hw:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    # ══════════════════════════════════════════════════════════════
    # V. RESULTS
    # ══════════════════════════════════════════════════════════════
    add_styled_paragraph(doc, "V. Results", 'Normal', font_size=12, bold=True,
                         space_before=12, space_after=6)

    add_styled_paragraph(doc, "A. Summary Comparison", 'Normal', font_size=11, bold=True,
                         italic=True, space_after=4)

    # Results Table
    results_table = doc.add_table(rows=4, cols=9)
    results_table.style = 'Table Grid'
    results_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    res_headers = ["Condition", "Temp", "Constr.", "Games", "LLM Calls",
                   "Legal", "Legal Rate", "Unique", "Latency (ms)"]
    for i, h in enumerate(res_headers):
        cell = results_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(8)
        set_cell_shading(cell, "D9E2F3")

    res_data = [
        ["t02_uncon_v2", "0.2", "No", "100", "211", "111", "52.6%", "7", "5,092"],
        ["t08_uncon", "0.8", "No", "100", "210", "110", "52.4%", "11", "11,676"],
        ["t08_constr", "0.8", "Yes", "20", "220", "220", "100.0%", "40", "10,084"],
    ]
    for row_idx, row_data in enumerate(res_data):
        for col_idx, val in enumerate(row_data):
            cell = results_table.rows[row_idx + 1].cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(8)
            # Highlight constrained row
            if row_idx == 2:
                set_cell_shading(cell, "E2EFDA")

    doc.add_paragraph()

    # B. Legal Move Rate
    add_styled_paragraph(doc, "B. Legal Move Rate Comparison", 'Normal', font_size=11, bold=True,
                         italic=True, space_after=4)

    legal_plot = os.path.join(PLOTS_DIR, "legal_rate_comparison.png")
    if os.path.isfile(legal_plot):
        doc.add_picture(legal_plot, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_styled_paragraph(doc, "Fig. 1. Legal Move Rate by Condition", 'Normal',
                             font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_after=6)

    add_styled_paragraph(
        doc,
        "Both unconstrained conditions cluster tightly around 52% (T=0.2: 52.6%, T=0.8: 52.4%), "
        "indicating a hard legality ceiling invariant to temperature. The constrained condition "
        "achieves a perfect 100.0% legal move rate across all 220 LLM calls.",
        'Normal', font_size=10, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6
    )

    # C. Pairwise Comparisons
    add_styled_paragraph(doc, "C. Pairwise Comparisons", 'Normal', font_size=11, bold=True,
                         italic=True, space_after=4)

    # Temperature effect table
    p = doc.add_paragraph(style='Normal')
    p.add_run("Temperature Effect (T=0.2 vs T=0.8, both unconstrained):").bold = True
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(4)

    temp_table = doc.add_table(rows=4, cols=2)
    temp_table.style = 'Table Grid'
    temp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    temp_data = [
        ["Metric", "Delta"],
        ["Legal rate", "−0.2 pp (52.6% → 52.4%)"],
        ["Unique moves", "+4 (7 → 11)"],
        ["Mean latency", "+6,583 ms"],
    ]
    for row_idx, row_data in enumerate(temp_data):
        for col_idx, val in enumerate(row_data):
            cell = temp_table.rows[row_idx].cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
            if row_idx == 0:
                set_cell_shading(cell, "D9E2F3")

    doc.add_paragraph()

    # Constraint effect table
    p = doc.add_paragraph(style='Normal')
    p.add_run("Constraint Effect (Unconstrained vs Constrained, both T=0.8):").bold = True
    p.runs[0].font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(4)

    constr_table = doc.add_table(rows=4, cols=2)
    constr_table.style = 'Table Grid'
    constr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    constr_data = [
        ["Metric", "Delta"],
        ["Legal rate", "+47.6 pp (52.4% → 100.0%)"],
        ["Unique moves", "+29 (11 → 40)"],
        ["Mean latency", "−1,592 ms"],
    ]
    for row_idx, row_data in enumerate(constr_data):
        for col_idx, val in enumerate(row_data):
            cell = constr_table.rows[row_idx].cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
            if row_idx == 0:
                set_cell_shading(cell, "D9E2F3")

    doc.add_paragraph()

    # D. Latency
    add_styled_paragraph(doc, "D. Response Latency Analysis", 'Normal', font_size=11, bold=True,
                         italic=True, space_after=4)

    latency_plot = os.path.join(PLOTS_DIR, "latency_comparison.png")
    if os.path.isfile(latency_plot):
        doc.add_picture(latency_plot, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_styled_paragraph(doc, "Fig. 2. Response Latency by Condition", 'Normal',
                             font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_after=6)

    latency_findings = [
        "T=0.2 unconstrained has the lowest mean latency (5,092 ms), consistent with reduced "
        "sampling effort at low temperature.",
        "T=0.8 unconstrained exhibits a high mean (11,676 ms) but low median (4,147 ms), "
        "indicating a heavily right-skewed distribution with occasional extreme outliers "
        "(max: 1,539,404 ms).",
        "T=0.8 constrained has a moderate mean (10,084 ms) and the highest median (9,980 ms), "
        "reflecting consistently longer inference times due to the larger prompt.",
    ]
    for finding in latency_findings:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(finding)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

    # E. Move Diversity
    add_styled_paragraph(doc, "E. Move Diversity Analysis", 'Normal', font_size=11, bold=True,
                         italic=True, space_before=8, space_after=4)

    diversity_plot = os.path.join(PLOTS_DIR, "move_diversity_comparison.png")
    if os.path.isfile(diversity_plot):
        doc.add_picture(diversity_plot, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_styled_paragraph(doc, "Fig. 3. Move Diversity (Unique Moves) by Condition", 'Normal',
                             font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             space_after=6)

    diversity_findings = [
        "T=0.2 unconstrained (7 unique moves): Low temperature leads to deterministic repetition. "
        "The model repeatedly produces g8f6, ignoring the evolving board state.",
        "T=0.8 unconstrained (11 unique moves): Higher temperature slightly increases diversity, "
        "but the model still converges on a narrow set of patterns.",
        "T=0.8 constrained (40 unique moves): Constrained decoding produces 3.6× more unique "
        "moves. By presenting the legal move list, the model selects from the full action space.",
    ]
    for finding in diversity_findings:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(finding)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

    # F. Game Completion
    add_styled_paragraph(doc, "F. Game Completion Analysis", 'Normal', font_size=11, bold=True,
                         italic=True, space_before=8, space_after=4)

    add_styled_paragraph(
        doc,
        "In the unconstrained conditions, all 200 games were aborted early due to illegal moves "
        "(early termination enabled), with most games lasting only 1–3 LLM calls before producing "
        "an illegal output. In the constrained condition, all 20 games ran to completion (reaching "
        "the 22-ply turn cap) with zero aborts, demonstrating sustained legal play over multiple turns.",
        'Normal', font_size=10, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6
    )

    # ══════════════════════════════════════════════════════════════
    # VI. DISCUSSION
    # ══════════════════════════════════════════════════════════════
    add_styled_paragraph(doc, "VI. Discussion", 'Normal', font_size=12, bold=True,
                         space_before=12, space_after=6)

    discussions = [
        ("A. The 52% Legality Ceiling",
         "The most striking finding is the remarkable consistency of the unconstrained legal move "
         "rate at approximately 52%, invariant to sampling temperature. This suggests that the "
         "model's legal move generation capability is limited by its internal representation of chess "
         "rules rather than by the stochasticity of the sampling process. The model appears to have "
         "learned a shallow statistical association between FEN-like strings and common UCI move "
         "patterns, but lacks a deep understanding of piece movement rules and board geometry."),

        ("B. Deterministic Repetition at Low Temperature",
         "The T=0.2 condition reveals a particularly informative failure mode: the model "
         "overwhelmingly produces the move g8f6 (knight to f6), a common opening response from "
         "Black. At near-greedy decoding, the model defaults to its highest-probability output "
         "token sequence regardless of the actual board state, effectively producing a \"cached\" "
         "response. This demonstrates that reducing temperature does not improve accuracy—it merely "
         "makes the model more confidently wrong."),

        ("C. Constraint Injection as a Hallucination Remedy",
         "The effectiveness of prompt-level constraint injection (100% legality) demonstrates that "
         "the model is capable of recognizing and selecting from a provided list of valid options, "
         "even when it cannot independently generate valid options. This is consistent with the "
         "distinction between generative and discriminative capabilities: the model fails at "
         "generation (producing a legal move from scratch) but succeeds at discrimination (selecting "
         "a legal move from a list). This has practical implications for LLM deployment in "
         "rule-governed domains: rather than expecting the model to internalize complex rule systems, "
         "one can externalize the rule enforcement into the prompt."),

        ("D. Diversity as a Side Effect of Constraints",
         "The increase in move diversity under constrained decoding (40 vs 11 unique moves) is a "
         "counterintuitive and valuable finding. One might expect that constraining the output space "
         "would reduce diversity, but the opposite occurs because: (1) without constraints, the "
         "model defaults to memorized high-frequency patterns; (2) with the legal move list "
         "presented, the model accesses the full action space and distributes selections more "
         "broadly. This suggests that constraint injection not only improves correctness but also "
         "combats the \"mode collapse\" behavior observed in unconstrained generation."),

        ("E. Limitations",
         "Results are specific to Llama 3.1 8B (Q4_0); other models may exhibit different legality "
         "ceilings. The 4-bit quantization may degrade chess reasoning capability. Latency "
         "measurements reflect CPU-only execution. With early termination, unconstrained runs "
         "provide limited data on deep game trees. We do not evaluate the strategic quality of "
         "moves (Elo rating), only their legality."),
    ]

    for title, text in discussions:
        add_styled_paragraph(doc, title, 'Normal', font_size=11, bold=True,
                             italic=True, space_after=4)
        add_styled_paragraph(doc, text, 'Normal', font_size=10,
                             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

    # ══════════════════════════════════════════════════════════════
    # VII. REPRODUCTION INSTRUCTIONS
    # ══════════════════════════════════════════════════════════════
    add_styled_paragraph(doc, "VII. Reproduction Instructions", 'Normal', font_size=12, bold=True,
                         space_before=12, space_after=6)

    add_styled_paragraph(doc, "The complete experiment can be reproduced in three steps:",
                         'Normal', font_size=10, space_after=6)

    add_styled_paragraph(doc, "Step 1: Environment Setup", 'Normal', font_size=10, bold=True,
                         space_after=2)
    code1 = "git clone https://github.com/Arpit-Panigrahi/llm-chess-engine.git\ncd llm-chess-engine\npip install -r requirements.txt\nollama pull llama3.1"
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(code1)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    add_styled_paragraph(doc, "Step 2: Run the Experiment Matrix", 'Normal', font_size=10,
                         bold=True, space_after=2)
    code2 = (
        "# Full matrix (3 conditions)\n"
        "bash scripts/run_experiment_matrix.sh \\\n"
        "  --model llama3.1 --num-games 100 --early-termination\n\n"
        "# Or run individual conditions:\n"
        "python3 scripts/run_game.py --temperature 0.2 \\\n"
        "  --no-constrained-decoding --num-games 100 \\\n"
        "  --early-termination --seed 42 --tag t02_unconstrained\n\n"
        "python3 scripts/run_game.py --temperature 0.8 \\\n"
        "  --no-constrained-decoding --num-games 100 \\\n"
        "  --early-termination --seed 42 --tag t08_unconstrained\n\n"
        "python3 scripts/run_game.py --temperature 0.8 \\\n"
        "  --constrained-decoding --num-games 20 \\\n"
        "  --max-turns 22 --seed 42 --tag t08_constrained"
    )
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(code2)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    add_styled_paragraph(doc, "Step 3: Generate Analysis Report", 'Normal', font_size=10,
                         bold=True, space_after=2)
    code3 = "python3 scripts/analyze_all.py \\\n  --run-root runs --out reports/experiment_matrix"
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(code3)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    # ══════════════════════════════════════════════════════════════
    # VIII. CONCLUSION
    # ══════════════════════════════════════════════════════════════
    add_styled_paragraph(doc, "VIII. Conclusion", 'Normal', font_size=12, bold=True,
                         space_before=12, space_after=6)

    add_styled_paragraph(
        doc,
        "This paper presents a controlled experimental study demonstrating that prompt-level legal "
        "move injection completely eliminates hallucinated chess moves in Llama 3.1 8B, raising the "
        "legal move rate from a temperature-invariant ceiling of approximately 52% to a perfect "
        "100%. Furthermore, constrained decoding increases move diversity by 3.6×, countering the "
        "mode collapse observed in unconstrained generation.",
        'Normal', font_size=10, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6
    )

    add_styled_paragraph(
        doc,
        "These findings suggest a practical design pattern for deploying LLMs in rule-governed "
        "domains: externalize rule enforcement into the prompt context, transforming the task from "
        "unconstrained generation to constrained selection. This approach requires no model "
        "modification, is compatible with any inference API, and can be applied to any domain where "
        "the set of valid outputs at each step can be enumerated.",
        'Normal', font_size=10, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6
    )

    add_styled_paragraph(
        doc,
        "Future work will extend this evaluation to additional models (GPT-4, Mistral, Qwen), "
        "evaluate the strategic quality of constrained move selections (Elo rating), and investigate "
        "hybrid architectures combining LLM-based evaluation with classical alpha-beta search.",
        'Normal', font_size=10, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12
    )

    # ══════════════════════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════════════════════
    doc.add_paragraph().add_run("─" * 80).font.size = Pt(6)

    add_styled_paragraph(doc, "References", 'Normal', font_size=12, bold=True,
                         space_before=8, space_after=6)

    references = [
        '[1] J. Wei et al., "Chain-of-thought prompting elicits reasoning in large language '
        'models," Advances in Neural Information Processing Systems, vol. 35, pp. 24824–24837, 2022.',

        '[2] Z. Ji et al., "Survey of hallucination in natural language generation," ACM Computing '
        'Surveys, vol. 55, no. 12, pp. 1–38, 2023.',

        '[3] S. Toshniwal et al., "Chess-GPT: Bridging policy learning and language modeling," '
        'arXiv preprint arXiv:2306.09200, 2023.',

        '[4] X. Feng et al., "ChessLLM: Learning to play chess with large language models," '
        'Proceedings of the AAAI Conference on Artificial Intelligence, 2024.',

        '[5] D. Silver et al., "Mastering the game of Go with deep neural networks and tree '
        'search," Nature, vol. 529, no. 7587, pp. 484–489, 2016.',

        '[6] R. Allbert, "VICE chess engine," Bluefever Software, 2013. [Online]. Available: '
        'https://github.com/bluefeversoft/vice',

        '[7] Ollama, "Ollama: Get up and running with large language models," 2024. [Online]. '
        'Available: https://ollama.com',

        '[8] N. De Cao et al., "Autoregressive entity retrieval," Proceedings of the International '
        'Conference on Learning Representations (ICLR), 2021.',
    ]

    for ref in references:
        p = doc.add_paragraph(style='Normal')
        run = p.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)

    doc.add_paragraph()
    add_styled_paragraph(
        doc,
        "Manuscript submitted July 2026. The complete source code, experiment data, and analysis "
        "pipeline are available at https://github.com/Arpit-Panigrahi/llm-chess-engine.",
        'Normal', font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6
    )

    # ── Save ──
    out_path = os.path.join(PROJECT_ROOT, "docs", "research_paper_ieee.docx")
    doc.save(out_path)
    print(f"✓ DOCX saved to: {out_path}")
    return out_path


if __name__ == "__main__":
    build_docx()
