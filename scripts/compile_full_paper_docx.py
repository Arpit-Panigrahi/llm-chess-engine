"""
compile_full_paper_docx.py
Compiles docs/research_paper_full.md → docs/research_paper_full.docx
with native styled Word tables and all 8 publication figures embedded.
"""
import os, re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

FIG_DIR = "reports/figures"
NAVY  = "1F497D"
GREEN = "1B5E20"
STRIPE= "F0F4FF"

def shd(cell, hex_col):
    ns = nsdecls("w")
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {ns} w:fill="{hex_col}"/>'))

def mar(cell, t=100, b=100, l=130, r=130):
    ns = nsdecls("w")
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:tcMar {ns}><w:top w:w="{t}" w:type="dxa"/>'
                  f'<w:bottom w:w="{b}" w:type="dxa"/>'
                  f'<w:left w:w="{l}" w:type="dxa"/>'
                  f'<w:right w:w="{r}" w:type="dxa"/></w:tcMar>'))

def style_table(table, hdr_bg=NAVY, stripe_bg=STRIPE):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for c_i, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                shd(cell, hdr_bg)
                mar(cell, 120, 120, 140, 140)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.bold = True; r.font.name = "Times New Roman"
                        r.font.size = Pt(9); r.font.color.rgb = RGBColor(255,255,255)
            else:
                if i % 2 == 0:
                    shd(cell, stripe_bg)
                mar(cell)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = "Times New Roman"; r.font.size = Pt(8.5)

def flush_md_table(doc, lines):
    """Parse markdown pipe table and emit a styled native Word table."""
    if not lines:
        return
    parsed = []
    for line in lines:
        if re.match(r'^\s*\|?\s*[-:]+[-| :]*$', line):
            continue
        cols = [c.strip() for c in line.strip().strip("|").split("|")]
        if any(cols):
            parsed.append(cols)
    if not parsed:
        return
    ncols = max(len(r) for r in parsed)
    for r in parsed:
        while len(r) < ncols:
            r.append("")

    tbl = doc.add_table(rows=len(parsed), cols=ncols)
    for ri, row in enumerate(parsed):
        for ci, val in enumerate(row):
            clean = re.sub(r'\*+', '', val).replace('`','').strip()
            p = tbl.cell(ri, ci).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
            rn = p.add_run(clean)
            rn.font.name = "Times New Roman"
            rn.font.size = Pt(9 if ri == 0 else 8.5)
            if ri == 0 or re.search(r'\*\*', val):
                rn.bold = True

    style_table(tbl)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)

def embed_figure(doc, img_path, caption):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        doc.add_picture(img_path, width=Inches(5.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        r = cap.add_run(caption)
        r.italic = True; r.font.size = Pt(9)
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(70, 70, 70)

def add_run_with_inline_fmt(p, text):
    """Split bold/italic markdown inline markers and add as runs."""
    parts = re.split(r'(\*\*.*?\*\*|\*[^*].*?[^*]\*(?!\*))', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1]); r.italic = True
        else:
            r = p.add_run(part)
        r.font.name = "Times New Roman"; r.font.size = Pt(10)

def build_docx():
    with open("docs/research_paper_full.md", encoding="utf-8") as f:
        lines = f.readlines()

    doc = docx.Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.1)
        sec.right_margin  = Inches(1.1)

    # ── Figure map keyed by "Figure N" ────────────────────────────────────────
    fig_map = {
        "Figure 1": (f"{FIG_DIR}/fig1_legality_rate.png",
                     "Figure 1: Legal Move Rate Across All Experimental Conditions (N=260 games, 1,077 inference calls)"),
        "Figure 2": (f"{FIG_DIR}/fig2_latency_decomposition.png",
                     "Figure 2: Granular Latency Decomposition — Cold-Start Disk Load vs. Warm Inference vs. p95 Tail Ceiling"),
        "Figure 3": (f"{FIG_DIR}/fig3_acpl_elo.png",
                     "Figure 3: Stockfish 18 Depth-12 ACPL & Estimated Elo (Regan–Guid Model). Lower ACPL = Higher Quality."),
        "Figure 4": (f"{FIG_DIR}/fig4_bpe_ablation.png",
                     "Figure 4: BPE Tokenizer Ablation Study — Representation Scheme vs. Legal Compliance Rate"),
        "Figure 5": (f"{FIG_DIR}/fig5_token_compression.png",
                     "Figure 5: Prompt Token Footprint — DMC reduces prefill compute by 45%"),
        "Figure 6": (f"{FIG_DIR}/fig6_cpl_distribution.png",
                     "Figure 6: Per-Move Centipawn Loss Distribution & Turn-by-Turn Trend — Fast Clamped DMC"),
        "Figure 7": (f"{FIG_DIR}/fig7_latency_scatter.png",
                     "Figure 7: Per-Call Inference Latency Profile — first call = cold-start, remaining = warm steady-state"),
        "Figure 8": (f"{FIG_DIR}/fig8_performance_profile.png",
                     "Figure 8: Multi-Dimensional Normalized Performance Profile across all conditions"),
    }

    in_code = False
    tbl_lines = []
    i = 0

    def flush_table():
        if tbl_lines:
            flush_md_table(doc, list(tbl_lines))
            tbl_lines.clear()

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip("\n")
        stripped = line.strip()
        i += 1

        # ── Code block toggle ─────────────────────────────────────────────────
        if stripped.startswith("```"):
            flush_table()
            in_code = not in_code
            continue

        if in_code:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run(line)
            r.font.name = "Courier New"; r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(50, 50, 50)
            continue

        # ── Markdown table rows ───────────────────────────────────────────────
        if "|" in stripped and not stripped.startswith("#"):
            tbl_lines.append(line)
            continue
        else:
            flush_table()

        if not stripped or stripped == "---":
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        if stripped.startswith("# "):
            flush_table()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(8)
            r = p.add_run(stripped[2:])
            r.bold = True; r.font.size = Pt(15)
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(31, 73, 125)

        elif stripped.startswith("## "):
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(5)
            r = p.add_run(stripped[3:])
            r.bold = True; r.font.size = Pt(13)
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(31, 73, 125)

        elif stripped.startswith("### "):
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            r = p.add_run(stripped[4:])
            r.bold = True; r.font.size = Pt(11)
            r.font.name = "Times New Roman"

        elif stripped.startswith("#### "):
            flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(stripped[5:])
            r.bold = True; r.italic = True; r.font.size = Pt(10)
            r.font.name = "Times New Roman"

        # ── Bold standalone table caption lines ───────────────────────────────
        elif stripped.startswith("**TABLE") or stripped.startswith("**Figure"):
            flush_table()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            clean = stripped.replace("**","")
            r = p.add_run(clean)
            r.bold = True; r.font.size = Pt(10)
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(31, 73, 125)

        # ── Bullet points ─────────────────────────────────────────────────────
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after   = Pt(2)
            p.paragraph_format.line_spacing  = 1.1
            add_run_with_inline_fmt(p, stripped[2:])

        # ── Numbered list ─────────────────────────────────────────────────────
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after   = Pt(3)
            p.paragraph_format.line_spacing  = 1.1
            text = re.sub(r'^\d+\.\s', '', stripped)
            add_run_with_inline_fmt(p, text)

        # ── Regular paragraph ─────────────────────────────────────────────────
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(5)
            p.paragraph_format.line_spacing = 1.15
            if any(kw in stripped for kw in ["Arpit","Vellore","Email","GitHub","School of"]):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run_with_inline_fmt(p, stripped)

        # ── Inline figure injection based on "(See Figure N)" ─────────────────
        for fig_key, (fig_path, fig_caption) in fig_map.items():
            if f"(See {fig_key})" in stripped:
                embed_figure(doc, fig_path, fig_caption)
                break

    flush_table()

    doc.save("docs/research_paper_full.docx")
    print("✓ docs/research_paper_full.docx compiled successfully")

if __name__ == "__main__":
    build_docx()
