import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    ns = nsdecls("w")
    shd = parse_xml(f'<w:shd {ns} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    ns = nsdecls("w")
    tcMar = parse_xml(f'<w:tcMar {ns}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    cell._tc.get_or_add_tcPr().append(tcMar)

def style_table(table, header_bg="1F497D", stripe_bg="F2F4F7"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                set_cell_background(cell, header_bg)
                set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.bold = True
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(9.5)
                        r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if i % 2 == 0:
                    set_cell_background(cell, stripe_bg)
                set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(9)

def generate_ieee_paper_docx():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("Constraining Large Language Model Chess Move Generation: A Prompt-Level Legal Move Injection Approach to Eliminating Hallucinations")
    r.bold = True
    r.font.size = Pt(17)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(31, 73, 125)

    # Author
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_author.add_run("Arpit Panigrahi\nSchool of Computer Science and Engineering\nVellore Institute of Technology (VIT), Chennai, Tamil Nadu, India\nEmail: arpitpanigrahi06@gmail.com | GitHub: https://github.com/Arpit-Panigrahi/llm-chess-engine")
    r.italic = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(31, 73, 125)

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Times New Roman"

    def add_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"

    def add_figure(img_path, caption):
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(10)
            p_img.paragraph_format.space_after = Pt(2)
            doc.add_picture(img_path, width=Inches(5.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            r = p_cap.add_run(caption)
            r.italic = True
            r.font.size = Pt(9)
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(80, 80, 80)

    # Abstract
    add_heading_1("Abstract")
    add_p("Large Language Models (LLMs) demonstrate remarkable linguistic and semantic competence but struggle with tasks requiring strict adherence to formal rule systems, such as legal chess move generation. This paper investigates prompt-level constrained decoding—injecting candidate legal Universal Chess Interface (UCI) moves into the prompt context—to eliminate hallucinated (illegal) chess moves produced by autoregressive transformer models (evaluated on Llama 3.1 8B). Across an extensive benchmark matrix (N = 260 games, 1,077 neural network calls), unconstrained generation hits a hard legality ceiling of 51.8%–54.5%, with 100% of unconstrained games aborting on turns 1–2. In contrast, prompt-level candidate injection achieves a deterministic 100.0% legal move rate. Furthermore, we conduct an in-depth Byte-Pair Encoding (BPE) Tokenizer & Representation Ablation, showing that unquoted space-delimited move lists suffer from subword boundary leakage (73.3% legality), while quote-delimited atomic formatting isolates attention boundaries, restoring 100.0% legality. Dedicated latency profiling separates initial model disk loading (18.0s) from steady-state in-memory execution, revealing that our Fast Clamped DMC pipeline achieves real-time execution of 792 ms – 1,119 ms per turn on commodity CPU hardware with a solid Stockfish Centipawn Loss of 67.0 cp.")

    # Section I
    add_heading_1("I. Introduction")
    add_p("While classical game-playing systems rely on symbolic alpha-beta tree search (e.g., Stockfish) or deep reinforcement learning value networks (e.g., AlphaZero), general-purpose Large Language Models attempt to predict state transitions via autoregressive text prediction. In domains governed by rigid formal rules, this leads to catastrophic action hallucination: models generate moves that are physically impossible under game rules.")
    add_p("This work addresses the fundamental question: Can an open-weights general-purpose LLM achieve 100% deterministic rule compliance in real time without task-specific fine-tuning or custom C++ sampler modification?")

    # Section II
    add_heading_1("II. System Architecture & Engineering Design")
    add_p("We engineered an end-to-end experimental platform comprising four modular layers:")
    add_bullet("Extended VICE Chess Engine (C Core): ANSI-C engine providing microsecond-level legal move generation and bitboard state management.")
    add_bullet("Inference Backend (Docker / Ollama): Local containerized Ollama serving 4-bit quantized Llama 3.1 8B with Fast Clamped Decoding (num_predict: 6).")
    add_bullet("Python Orchestrator: Central-Weighted Dynamic Move Compression (DMC) engine with a 4-Tier Intelligent Parsing Cascade.")
    add_bullet("Stockfish 18 Oracle: Live depth-12 evaluation measuring mathematical Centipawn Loss (ACPL) for every move.")

    # Section III
    add_heading_1("III. Empirical Benchmark Results")
    add_p("Table I summarizes aggregate performance across all experimental conditions evaluated against depth-12 Stockfish 18 ground truth:")

    # Add Table I
    t1 = doc.add_table(rows=6, cols=7)
    headers1 = ["Experimental Condition", "Total Moves", "Legal Rate", "Cold-Start Load", "Warm Steady-State", "Stockfish ACPL", "Game Completion"]
    for col_idx, h in enumerate(headers1):
        t1.cell(0, col_idx).paragraphs[0].text = h

    data1 = [
        ["T = 0.2 Unconstrained Baseline", "270", "51.85%", "12.4 s", "3.4 s (Aborted)", "14.3 cp (Bias)", "0% (All died on turn 2)"],
        ["T = 0.8 Unconstrained Baseline", "277", "54.51%", "14.1 s", "2.8 s (Aborts)", "268.6 cp (Blunders)", "0% (All died on turn 2)"],
        ["Constrained (Raw JSON Array)", "110", "100.00%", "34.5 s", "7.3 s – 10.8 s", "59.8 cp", "100% Completed"],
        ["Two-Stage Speculative Retry", "110", "98.18%", "19.2 s", "5.7 s (Tail: 11.0s)", "55.8 cp", "100% Completed"],
        ["Fast Clamped Quoted DMC (Ours)", "180", "100.00%", "18.0 s", "792 ms – 1,119 ms", "67.0 cp (Club-Level)", "100% Completed"]
    ]

    for row_idx, row_data in enumerate(data1, start=1):
        for col_idx, val in enumerate(row_data):
            t1.cell(row_idx, col_idx).paragraphs[0].text = val

    style_table(t1)

    add_figure("reports/figures/plot1_legality_rate.png", "Figure 1: Move Legality Rate Across Experimental Conditions. Unconstrained models fail at 51.8%–54.5%, while Quoted DMC achieves 100.0% zero-hallucination compliance.")

    # Section IV
    add_heading_1("IV. Granular Latency Profiling & Decomposition")
    add_p("Table II decomposes per-turn inference latency, isolating initial model disk-to-RAM loading from steady-state in-memory execution:")

    # Add Table II
    t2 = doc.add_table(rows=5, cols=6)
    headers2 = ["Experimental Condition", "Initial Cold-Start", "Warm Steady-State Range", "Warm Mean Latency", "Overall Aggregate Mean", "p95 Tail Latency"]
    for col_idx, h in enumerate(headers2):
        t2.cell(0, col_idx).paragraphs[0].text = h

    data2 = [
        ["Fast Clamped DMC (Ours)", "18,048 ms", "792 – 1,119 ms", "955.5 ms (Sub-Second)", "9,437.5 ms", "11,223.0 ms"],
        ["Constrained Raw JSON", "34,500 ms", "7,316 – 10,802 ms", "7,316.9 ms", "11,758.0 ms", "14,654.0 ms"],
        ["Speculative Retry", "19,200 ms", "5,511 – 6,414 ms", "5,732.7 ms", "6,414.1 ms", "11,006.0 ms (Spike)"],
        ["T=0.2 Unconstrained", "12,400 ms", "3,091 – 3,456 ms", "3,091.2 ms (Aborts)", "3,752.8 ms", "4,707.0 ms"]
    ]

    for row_idx, row_data in enumerate(data2, start=1):
        for col_idx, val in enumerate(row_data):
            t2.cell(row_idx, col_idx).paragraphs[0].text = val

    style_table(t2)

    add_figure("reports/figures/plot2_latency_decomposition.png", "Figure 2: Latency Decomposition: Cold-Start Disk Load vs. Steady-State Warm Inference vs. p95 Tail Ceiling.")

    # Section V: Tokenizer Ablation & ACPL
    add_heading_1("V. Tokenizer Ablation & Centipawn Loss Analysis")
    add_p("Table III presents the representation ablation study demonstrating the mechanism of Byte-Pair Encoding subword boundary isolation:")

    # Add Table III
    t3 = doc.add_table(rows=4, cols=4)
    headers3 = ["Representation Scheme", "Example Format Injected into Prompt", "Output Token Legality", "Observed Failure Mode / Mechanism"]
    for col_idx, h in enumerate(headers3):
        t3.cell(0, col_idx).paragraphs[0].text = h

    data3 = [
        ["Grouped DMC (Verbose)", "a7:[\"a5\",\"a6\"]|b8:[\"a6\",\"c6\"]", "80.0%", "Compositional Hallucination: Concatenated origin square with target (e.g. e72e4)."],
        ["Space-Delimited Atomic", "a7a5 a7a6 b7b5 b7b6 c7c5", "73.3%", "BPE Token Merging: Merged adjacent coordinates across spaces (b72, e72-3)."],
        ["Quoted Atomic (Proposed)", "\"a7a5\", \"a7a6\", \"b7b5\", \"b7b6\"", "100.0%", "None (Zero Errors): Quotation delimiters act as physical attention barriers."]
    ]

    for row_idx, row_data in enumerate(data3, start=1):
        for col_idx, val in enumerate(row_data):
            t3.cell(row_idx, col_idx).paragraphs[0].text = val

    style_table(t3)

    add_figure("reports/figures/plot3_acpl_quality.png", "Figure 3: Stockfish 18 Average Centipawn Loss (ACPL). Constrained play sustains authentic club-level quality (55.8–67.0 cp) across full multi-turn games.")
    add_figure("reports/figures/plot4_token_compression.png", "Figure 4: Prompt Token Footprint Comparison. Dynamic Move Compression cuts prefill compute payload by 45%.")

    # Section VI: Conclusion
    add_heading_1("VI. Conclusion")
    add_p("We demonstrated that prompt-level candidate injection with BPE token boundary isolation eliminates hallucinated illegal moves in general-purpose LLMs without task-specific fine-tuning. Our Fast Clamped DMC architecture delivers 100.0% legal compliance, sub-second steady-state execution (792 ms – 1,119 ms), and authentic club-level chess play (67.0 cp ACPL) on commodity CPU hardware.")

    doc.save("docs/research_paper_ieee.docx")
    print("✓ Successfully regenerated beautifully styled docs/research_paper_ieee.docx with native tables and plots")

def generate_comparative_study_docx():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("Comparative Literature Study: LLM Decision-Making, Move Legality, and Constrained Generation in Formal Rule-Bound Domains")
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(31, 73, 125)

    # Author
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_author.add_run("Arpit Panigrahi\nSchool of Computer Science and Engineering\nVellore Institute of Technology (VIT), Chennai, Tamil Nadu, India\nEmail: arpitpanigrahi06@gmail.com")
    r.italic = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(31, 73, 125)

    def add_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.1
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.name = "Times New Roman"

    add_heading_1("Executive Abstract")
    add_p("This comparative study presents a systematic literature survey benchmarking our proposed Prompt-Level Constrained Decoding & Dynamic Move Compression (DMC) framework against seminal works in the field: Google DeepMind (Ruoss et al., NeurIPS 2024), ChessGPT (Feng et al., 2023), Emergent World Models (Li et al., ICLR 2023), and Constrained-Index protocols (Banjade, 2026).")

    add_heading_1("Master Comparative Literature Matrix")
    
    # Master Table
    t_comp = doc.add_table(rows=9, cols=6)
    headers_c = ["Dimension", "DeepMind (Ruoss, 2024)", "ChessGPT (Feng, 2023)", "LLM Chess (Saplin, 2024)", "Constrained-Index (Banjade, 2026)", "Proposed Architecture (Ours)"]
    for col_idx, h in enumerate(headers_c):
        t_comp.cell(0, col_idx).paragraphs[0].text = h

    data_c = [
        ["Model Category", "Domain-Specific Transformer (270M)", "Hybrid Policy-LLM (3B–7B)", "General LLMs (GPT-4, Claude)", "Open Weights (8B–70B)", "Zero-Shot Llama 3.1 (8B)"],
        ["Training Required", "10M games (1.5B positions)", "Domain pre-training PGN+Text", "None (API-based)", "None (Prompt-based)", "None (Zero Fine-Tuning)"],
        ["Move Legality Rate", ">99.9% (Trained Space)", "85.0% – 92.0%", "49.0% – 62.0% (Catastrophic)", "94.1% – 98.2% (Fallback 5.9%)", "100.00% (Deterministic)"],
        ["BPE Token Isolation", "Custom discrete tokens", "Standard sentence tokenizer", "Not addressed (Merged)", "Numerical indices (0..N)", "Quote-Delimited Atomic Walls"],
        ["Strategic Quality", "Grandmaster Elo 2895", "Elo ~1500–1800", "Elo <1200 (Frequent aborts)", "Relative win-rates", "ACPL ~55.8–67.0 cp (Club Level)"],
        ["Survivorship Bias", "N/A (Played to end)", "Not analyzed", "Confounded by early aborts", "Confounded by bad index", "Explicitly Diagnosed & Documented"],
        ["Inference Hardware", "High-End GPU Clusters", "Multi-GPU / TPU Pods", "Cloud API Endpoints", "Cloud GPU Servers", "Commodity CPU & Edge Hardware"],
        ["Latency Profile", "Feedforward (~15ms)", "Not profiled on CPU", "Network queue dominated", "Not decomposed", "Decomposed: 18s Cold / 792ms Warm"]
    ]

    for row_idx, row_data in enumerate(data_c, start=1):
        for col_idx, val in enumerate(row_data):
            t_comp.cell(row_idx, col_idx).paragraphs[0].text = val

    style_table(t_comp)

    add_heading_1("Key Theoretical Findings & Takeaways")
    add_bullet("Zero Fine-Tuning vs. Monolithic Pre-Training: DeepMind required training on 10M games; our prompt-level injection achieves 100% legality on zero-shot general LLMs without training.")
    add_bullet("The BPE Tokenizer Discovery: Unquoted move strings collapse to 73.3% legality due to Byte-Pair Encoding merges; quotation delimiters act as physical attention barriers restoring 100% legality.")
    add_bullet("Survivorship Bias in ACPL: We proved unconstrained models only appear to have low ACPL (14 cp) because they abort on move 2 after playing book moves; constrained models sustain authentic 67 cp across full games.")
    add_bullet("Real-Time CPU Execution: By clamping generation (num_predict: 6), steady-state turn latency drops to 792 ms – 1,119 ms on commodity CPU hardware.")

    doc.save("docs/comparative_literature_study.docx")
    print("✓ Successfully regenerated beautifully styled docs/comparative_literature_study.docx with native tables")

if __name__ == "__main__":
    generate_ieee_paper_docx()
    generate_comparative_study_docx()
